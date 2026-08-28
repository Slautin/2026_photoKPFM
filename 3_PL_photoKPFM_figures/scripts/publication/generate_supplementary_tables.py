"""Build manuscript-numbered PhotoKPFM supplementary Tables S1-S4."""

from __future__ import annotations

import argparse
import hashlib
import json
import math
import os
import tempfile
from datetime import UTC, datetime
from pathlib import Path

import h5py
import numpy as np
import pandas as pd
from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(__file__).resolve().parents[2]
KPFM_H5 = ROOT / "data" / "raw" / "photokpfm_measurements.h5"
PLATE_MAP_CSV = ROOT / "data" / "processed" / "plate_composition_map.csv"
SUMMARY_CSV = ROOT / "data" / "literature" / "table_s3_relationship_summary.csv"
RELATIONSHIPS_CSV = ROOT / "data" / "literature" / "table_s4_retained_relationships.csv"
ASSOCIATIONS_CSV = ROOT / "data" / "processed" / "exploratory_association_statistics.csv"
TABLE_S1_CSV = ROOT / "data" / "processed" / "photokpfm_acquisition_log.csv"
DEFAULT_OUTPUT = ROOT / "results" / "tables" / "supplementary_tables_current_manuscript_S1-S4.docx"

TABLE_S1_CAPTION = (
    "Table S1. Ternary-library photoKPFM experimental acquisition log. Each row is one paired dark and "
    "illuminated acquisition stored in the supplied raw HDF5 file. Plate indices are zero-based and are "
    "mapped to the included 96-well nominal-composition lookup. Ternary scans contain 128 × 128 pixels over "
    "a 5 × 5 µm² area, as stated in the current manuscript Methods. The HDF5 file contains acquisition-order "
    "identifiers but no exact clock timestamps or external calibration metadata; neither is inferred here. "
    "Surface-potential measurements remain relative, uncalibrated instrument signals."
)

TABLE_S2_CAPTION = (
    "Table S2. Evidence-grounded scientific basis for complementary PL and KPFM measurements. "
    "KPFM-derived contact-potential, surface-potential, and work-function observables describe "
    "electrostatic or energetic behavior, whereas steady-state PL intensity and PLQY describe "
    "emissive output and TRPL or carrier lifetime describes carrier kinetics. Relationships involving "
    "steady-state PL intensity are most directly relevant to the present PL peak-intensity comparison. "
    "PLQY relationships provide related optical context, whereas TRPL and carrier-lifetime relationships "
    "provide broader carrier-dynamics context because those measurements were not performed in the present "
    "study. Counts report 15 retained paired relationships from seven papers. Co-observation supports a "
    "complementary interpretation and does not establish a causal KPFM-to-PL pathway."
)

TABLE_S3_CAPTION = (
    "Table S3. Evidence-level records for retained KPFM–PL-family relationships. The table reports "
    "15 retained relationships with the sample or intervention context, paired KPFM and PL-family "
    "observables, direct reported results, evidence-grounded claim, complementary interpretation, "
    "interpretation limit, evidence strength, and causal-status classification. Relationships involving "
    "steady-state PL intensity or PLQY are most relevant to the present optical comparison, whereas TRPL "
    "and carrier-lifetime relationships provide broader context. All retained relationships are classified "
    "as complementary-only and do not establish a causal KPFM-to-PL pathway."
)

TABLE_S4_CAPTION = (
    "Table S4. Exploratory raw and composition-adjusted PL–photoKPFM association tests. The table reports "
    "Pearson and Spearman coefficients, bootstrap confidence intervals, leave-one-out sensitivity, partial "
    "correlations adjusted for nominal PEA and BDA precursor fractions, and Benjamini-Hochberg q-values for "
    "36 exploratory comparisons. None survived multiplicity correction at q < 0.05. The analyses are "
    "descriptive and do not support a causal relationship between the PL and photoKPFM observables. "
    "PL-derived fitted emissive components, where included, are spectral-model outputs and should not be "
    "interpreted as structurally confirmed phase fractions."
)

S2_COLUMNS = {
    "kpfm_observable_class": "KPFM observable class",
    "pl_family_observable_class": "PL-family observable class",
    "retained_relationships": "Retained relationships",
    "evidence_grounded_scientific_claim": "Evidence-grounded scientific claim",
    "how_the_measurements_are_complementary": "How the measurements are complementary",
    "relevance_to_the_present_experiment": "Relevance to the present experiment",
    "interpretation_limit": "Interpretation limit",
}

S3_COLUMNS = {
    "relationship_id": "Relationship ID",
    "paper_title": "Paper title",
    "doi": "DOI",
    "intervention_or_comparison": "Intervention or comparison",
    "kpfm_observable_and_direct_result": "KPFM observable and direct result",
    "pl_family_observable_and_direct_result": "PL-family observable and direct result",
    "evidence_grounded_claim": "Evidence-grounded claim",
    "complementary_interpretation": "Complementary interpretation",
    "interpretation_limit": "Interpretation limit",
    "evidence_strength": "Evidence strength",
    "causal_status": "Causal status",
}

KPFM_LABELS = {
    "photovoltage_abs_V": "|ΔSP|",
    "photovoltage_signed_V": "Signed ΔSP (light - dark)",
    "mu_dark_V": "Dark relative surface potential",
    "mu_light_V": "Illuminated relative surface potential",
}

PL_LABELS = {
    "Top_PL_peak_intensity": "Top-read PL peak intensity",
    "dominant_peak_nm": "Dominant PL peak wavelength",
    "frac_n=1-like / 2D": "PL-fitted n=1-like/2D emissive fraction",
    "frac_n=2-like": "PL-fitted n=2-like emissive fraction",
    "frac_n=3-like": "PL-fitted n=3-like emissive fraction",
    "frac_n=4-like": "PL-fitted n=4-like emissive fraction",
    "frac_n>=5 / high-n": "PL-fitted n≥5/high-n emissive fraction",
    "frac_3D / 3D-like": "PL-fitted 3D/3D-like emissive fraction",
    "frac_low_n_1_to_4": "PL-fitted low-n (n=1-4) emissive fraction",
}

PLATE_ROWS = "ABCDEFGH"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--output", type=Path, default=DEFAULT_OUTPUT)
    parser.add_argument("--overwrite", action="store_true")
    return parser.parse_args()


def sha256(path: Path) -> str:
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def relative(path: Path) -> str:
    return path.resolve().relative_to(ROOT.resolve()).as_posix()


def clean_text(value: object) -> str:
    if pd.isna(value):
        return ""
    return str(value).strip()


def format_coefficient(value: object) -> str:
    if pd.isna(value):
        return "NA (not estimable)"
    return f"{float(value):.3f}".replace("-", "−")


def format_probability(value: object) -> str:
    if pd.isna(value):
        return "NA (not estimable)"
    number = float(value)
    return "<0.001" if number < 0.001 else f"{number:.3f}"


def format_interval(low: object, high: object) -> str:
    if pd.isna(low) or pd.isna(high):
        return "NA (not estimable)"
    return f"[{format_coefficient(low)}, {format_coefficient(high)}]"


def require_columns(frame: pd.DataFrame, expected: set[str], label: str) -> None:
    missing = sorted(expected.difference(frame.columns))
    if missing:
        raise ValueError(f"{label} is missing required columns: {missing}")


def decode_identifier(value: object) -> str:
    if isinstance(value, bytes):
        return value.decode("utf-8", errors="strict")
    return str(value)


def plate_index_to_well(index: int) -> str:
    if not 0 <= index < 96:
        raise ValueError(f"Invalid zero-based plate index: {index}")
    return f"{PLATE_ROWS[index % 8]}{index // 8 + 1}"


def build_table_s1() -> pd.DataFrame:
    plate = pd.read_csv(PLATE_MAP_CSV)
    plate_columns = {"Well", "FA_pct", "BDA_pct", "PEA_pct"}
    require_columns(plate, plate_columns, "Table S1 plate map")
    if len(plate) != 96 or plate["Well"].duplicated().any():
        raise ValueError("Table S1 plate map must contain 96 unique wells.")
    totals = plate[["FA_pct", "BDA_pct", "PEA_pct"]].apply(pd.to_numeric, errors="raise").sum(axis=1)
    if not totals.eq(100).all():
        raise ValueError("Table S1 plate-map compositions must sum to 100 percent.")
    lookup = plate.set_index("Well")
    with h5py.File(KPFM_H5, "r") as handle:
        required = {"idx", "dark_fn", "light_fn", "dark_data", "light_data", "X", "X_train", "y_train"}
        missing = sorted(required.difference(handle.keys()))
        if missing:
            raise ValueError(f"Table S1 HDF5 source is missing datasets: {missing}")
        indices = [int(value) for value in handle["idx"][...]]
        dark_ids = [decode_identifier(value) for value in handle["dark_fn"][...]]
        light_ids = [decode_identifier(value) for value in handle["light_fn"][...]]
        dark_shape = tuple(handle["dark_data"].shape)
        light_shape = tuple(handle["light_data"].shape)
        grid_coordinates = np.asarray(handle["X"], dtype=float)
        training_coordinates = np.asarray(handle["X_train"], dtype=float)
        training_rows = int(handle["X_train"].shape[0])
        response_rows = int(handle["y_train"].shape[0])
    lengths = {len(indices), len(dark_ids), len(light_ids), training_rows, response_rows}
    if lengths != {19}:
        raise ValueError(f"Table S1 HDF5 acquisition arrays must each contain 19 rows; found {sorted(lengths)}.")
    if dark_shape != (19, 128, 128) or light_shape != (19, 128, 128):
        raise ValueError(
            f"Table S1 expected paired 19 × 128 × 128 image arrays; found {dark_shape} and {light_shape}."
        )
    if not np.allclose(training_coordinates, grid_coordinates[indices]):
        raise ValueError("Table S1 HDF5 training coordinates do not match the indexed composition grid.")
    rows: list[dict[str, object]] = []
    for iteration, (plate_index, dark_id, light_id) in enumerate(
        zip(indices, dark_ids, light_ids), start=1
    ):
        well = plate_index_to_well(plate_index)
        if well not in lookup.index:
            raise ValueError(f"Table S1 plate index {plate_index} maps to missing well {well}.")
        composition = lookup.loc[well]
        expected_coordinate = np.asarray(
            [
                float(composition["PEA_pct"]) / 100.0 + 0.5 * float(composition["FA_pct"]) / 100.0,
                (math.sqrt(3.0) / 2.0) * float(composition["FA_pct"]) / 100.0,
            ]
        )
        if not np.allclose(grid_coordinates[plate_index], expected_coordinate):
            raise ValueError(
                f"Table S1 HDF5 index {plate_index} is inconsistent with the nominal composition for {well}."
            )
        rows.append(
            {
                "Acquisition iteration": iteration,
                "Plate index (zero-based)": plate_index,
                "Well ID": well,
                "Nominal FAPbI3 (%)": int(composition["FA_pct"]),
                "Nominal BDAPbI4 (%)": int(composition["BDA_pct"]),
                "Nominal PEA2PbI4 (%)": int(composition["PEA_pct"]),
                "Dark acquisition ID": dark_id,
                "Illuminated acquisition ID": light_id,
                "Scan pixels": "128 × 128",
                "Scan area": "5 × 5 µm²",
            }
        )
    result = pd.DataFrame(rows).map(clean_text)
    if result["Well ID"].nunique() != 13:
        raise ValueError(f"Table S1 must contain 13 unique wells; found {result['Well ID'].nunique()}.")
    return result


def build_table_s2() -> pd.DataFrame:
    source = pd.read_csv(SUMMARY_CSV)
    require_columns(source, set(S2_COLUMNS), "Table S2 source")
    if len(source) != 6:
        raise ValueError(f"Table S2 requires six observable-pair rows; found {len(source)}.")
    counts = pd.to_numeric(source["retained_relationships"], errors="raise")
    if int(counts.sum()) != 15:
        raise ValueError(f"Table S2 retained counts must sum to 15; found {int(counts.sum())}.")
    result = source[list(S2_COLUMNS)].rename(columns=S2_COLUMNS).copy()
    return result.map(clean_text)


def build_table_s3() -> pd.DataFrame:
    source = pd.read_csv(RELATIONSHIPS_CSV)
    require_columns(source, set(S3_COLUMNS), "Table S3 source")
    if len(source) != 15:
        raise ValueError(f"Table S3 requires 15 retained relationships; found {len(source)}.")
    if source["relationship_id"].duplicated().any():
        raise ValueError("Table S3 contains duplicate relationship identifiers.")
    if "R08" in set(source["relationship_id"].astype(str)):
        raise ValueError("Removed relationship R08 remains in the Table S3 source.")
    statuses = sorted(source["causal_status"].dropna().astype(str).unique().tolist())
    if statuses != ["complementary_only"]:
        raise ValueError(f"Unexpected Table S3 causal-status values: {statuses}")
    result = source[list(S3_COLUMNS)].rename(columns=S3_COLUMNS).copy()
    return result.map(clean_text)


def build_table_s4() -> pd.DataFrame:
    source = pd.read_csv(ASSOCIATIONS_CSV)
    required = {
        "kpfm_metric",
        "pl_metric",
        "n",
        "pearson_r",
        "pearson_p",
        "pearson_ci_low",
        "pearson_ci_high",
        "spearman_rho",
        "spearman_p",
        "spearman_loo_low",
        "spearman_loo_high",
        "partial_spearman_rho",
        "partial_spearman_p",
        "spearman_q_bh",
        "partial_spearman_q_bh",
        "causal_claim_supported",
    }
    require_columns(source, required, "Table S4 source")
    if len(source) != 36:
        raise ValueError(f"Table S4 requires 36 exploratory tests; found {len(source)}.")
    if set(source["kpfm_metric"].astype(str)) != set(KPFM_LABELS):
        raise ValueError("Table S4 contains an unexpected KPFM metric set.")
    if set(source["pl_metric"].astype(str)) != set(PL_LABELS):
        raise ValueError("Table S4 contains an unexpected PL metric set.")
    if (pd.to_numeric(source["spearman_q_bh"], errors="raise") < 0.05).any():
        raise ValueError("A raw Table S4 association unexpectedly survived BH correction.")
    if (pd.to_numeric(source["partial_spearman_q_bh"], errors="raise") < 0.05).any():
        raise ValueError("An adjusted Table S4 association unexpectedly survived BH correction.")
    causal = source["causal_claim_supported"].astype(str).str.casefold()
    if causal.isin({"true", "1", "yes"}).any():
        raise ValueError("A Table S4 row unexpectedly supports a causal claim.")
    primary = source["kpfm_metric"].eq("photovoltage_abs_V") & source["pl_metric"].eq(
        "Top_PL_peak_intensity"
    )
    if int(primary.sum()) != 1:
        raise ValueError("The primary |ΔSP|-versus-PL comparison is not unique.")
    ordered = pd.concat(
        [
            source.loc[primary],
            source.loc[~primary].sort_values(["kpfm_metric", "pl_metric"]),
        ],
        ignore_index=True,
    )
    rows: list[dict[str, object]] = []
    for row in ordered.to_dict(orient="records"):
        rows.append(
            {
                "KPFM metric": KPFM_LABELS[str(row["kpfm_metric"])],
                "PL metric": PL_LABELS[str(row["pl_metric"])],
                "n": int(row["n"]),
                "Pearson r": format_coefficient(row["pearson_r"]),
                "Pearson p": format_probability(row["pearson_p"]),
                "Pearson bootstrap 95% CI": format_interval(
                    row["pearson_ci_low"], row["pearson_ci_high"]
                ),
                "Spearman ρ": format_coefficient(row["spearman_rho"]),
                "Spearman p": format_probability(row["spearman_p"]),
                "Spearman BH q": format_probability(row["spearman_q_bh"]),
                "Composition controls": "Nominal PEA and BDA precursor fractions",
                "Partial Spearman ρ": format_coefficient(row["partial_spearman_rho"]),
                "Partial Spearman p": format_probability(row["partial_spearman_p"]),
                "Partial Spearman BH q": format_probability(row["partial_spearman_q_bh"]),
                "Leave-one-out Spearman range": format_interval(
                    row["spearman_loo_low"], row["spearman_loo_high"]
                ),
                "Interpretation status": "Exploratory; no causal claim supported",
            }
        )
    return pd.DataFrame(rows).map(clean_text)


def set_repeat_header(row) -> None:
    properties = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    properties.append(header)


def set_cell_width(cell, width_inches: float) -> None:
    width = Inches(width_inches)
    cell.width = width
    properties = cell._tc.get_or_add_tcPr()
    element = properties.find(qn("w:tcW"))
    if element is None:
        element = OxmlElement("w:tcW")
        properties.append(element)
    element.set(qn("w:w"), str(int(width)))
    element.set(qn("w:type"), "dxa")


def format_run(run, size: float, bold: bool = False) -> None:
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), "Times New Roman")


def add_caption(document: Document, text: str) -> None:
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    label, separator, body = text.partition(".")
    format_run(paragraph.add_run(label + separator), 8.5, bold=True)
    if body:
        format_run(paragraph.add_run(body), 8.5)


def add_dataframe_table(
    document: Document,
    frame: pd.DataFrame,
    widths: list[float],
    font_size: float,
) -> None:
    if len(widths) != len(frame.columns):
        raise ValueError("Column-width count does not match table columns.")
    table = document.add_table(rows=1, cols=len(frame.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_repeat_header(table.rows[0])
    for index, (column, width) in enumerate(zip(frame.columns, widths)):
        cell = table.rows[0].cells[index]
        set_cell_width(cell, width)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraph.paragraph_format.space_after = Pt(0)
        format_run(paragraph.add_run(str(column)), font_size, bold=True)
    for values in frame.itertuples(index=False, name=None):
        row = table.add_row()
        for index, (value, width) in enumerate(zip(values, widths)):
            cell = row.cells[index]
            set_cell_width(cell, width)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.0
            format_run(paragraph.add_run(str(value)), font_size)


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Inches(14)
    section.page_height = Inches(8.5)
    section.top_margin = Inches(0.35)
    section.bottom_margin = Inches(0.35)
    section.left_margin = Inches(0.35)
    section.right_margin = Inches(0.35)
    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(8.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    properties = document.core_properties
    properties.title = "PhotoKPFM/PL supplementary Tables S1-S4"
    properties.subject = "Manuscript-numbered supplementary tables generated from release CSV sources"
    properties.author = "PhotoKPFM/PL manuscript release"
    properties.keywords = "PhotoKPFM, photoluminescence, supplementary tables"


def build_document(
    s1: pd.DataFrame,
    s2: pd.DataFrame,
    s3: pd.DataFrame,
    s4: pd.DataFrame,
) -> Document:
    document = Document()
    configure_document(document)
    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.paragraph_format.space_after = Pt(6)
    format_run(title.add_run("PhotoKPFM/PL Supplementary Tables S1-S4"), 12, bold=True)
    add_caption(document, TABLE_S1_CAPTION)
    add_dataframe_table(
        document,
        s1,
        [0.75, 0.85, 0.55, 0.85, 0.85, 0.85, 1.05, 1.15, 0.75, 0.8],
        7.0,
    )
    document.add_page_break()
    add_caption(document, TABLE_S2_CAPTION)
    add_dataframe_table(document, s2, [1.2, 1.2, 0.75, 2.35, 2.25, 2.4, 2.5], 7.0)
    document.add_page_break()
    add_caption(document, TABLE_S3_CAPTION)
    add_dataframe_table(
        document,
        s3,
        [0.45, 1.45, 0.9, 1.05, 1.4, 1.5, 1.35, 1.25, 1.35, 0.7, 0.75],
        6.1,
    )
    document.add_page_break()
    add_caption(document, TABLE_S4_CAPTION)
    add_dataframe_table(
        document,
        s4,
        [0.9, 1.3, 0.35, 0.55, 0.55, 1.15, 0.6, 0.55, 0.6, 1.15, 0.7, 0.6, 0.7, 1.15, 1.25],
        6.2,
    )
    return document


def validate_saved_document(path: Path, expected: list[pd.DataFrame]) -> None:
    document = Document(path)
    if len(document.tables) != 4:
        raise ValueError(f"Expected four generated tables; found {len(document.tables)}.")
    if [len(table.rows) for table in document.tables] != [len(frame) + 1 for frame in expected]:
        raise ValueError("Generated Word table row counts do not match source data.")
    text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    for label in ("Table S1.", "Table S2.", "Table S3.", "Table S4."):
        if label not in text:
            raise ValueError(f"Generated document is missing {label}")
    for table, frame in zip(document.tables, expected):
        observed = [[cell.text for cell in row.cells] for row in table.rows]
        wanted = [list(map(str, frame.columns))] + [list(map(str, row)) for row in frame.itertuples(index=False, name=None)]
        if observed != wanted:
            raise ValueError("Generated Word table values differ from the prepared release data.")


def write_outputs(output: Path, overwrite: bool) -> tuple[Path, Path, Path]:
    provenance = output.with_suffix(".provenance.json")
    conflicts = [path for path in (output, provenance, TABLE_S1_CSV) if path.exists()]
    if conflicts and not overwrite:
        names = ", ".join(str(path) for path in conflicts)
        raise FileExistsError(f"Output exists: {names}. Use --overwrite to replace it.")
    sources = (KPFM_H5, PLATE_MAP_CSV, SUMMARY_CSV, RELATIONSHIPS_CSV, ASSOCIATIONS_CSV)
    for source in sources:
        if not source.is_file():
            raise FileNotFoundError(source)
    s1 = build_table_s1()
    s2 = build_table_s2()
    s3 = build_table_s3()
    s4 = build_table_s4()
    output.parent.mkdir(parents=True, exist_ok=True)
    with tempfile.TemporaryDirectory(prefix=".photokpfm_tables_", dir=output.parent) as directory:
        temporary = Path(directory) / output.name
        build_document(s1, s2, s3, s4).save(temporary)
        validate_saved_document(temporary, [s1, s2, s3, s4])
        temporary_s1_csv = Path(directory) / TABLE_S1_CSV.name
        s1.to_csv(temporary_s1_csv, index=False, encoding="utf-8")
        if not pd.read_csv(temporary_s1_csv, dtype=str).fillna("").equals(s1.astype(str)):
            raise ValueError("Generated Table S1 CSV differs from the prepared acquisition log.")
        output_record = {
            "status": "complete",
            "generated_at_utc": datetime.now(UTC).isoformat(),
            "table_s1": {
                "generated": True,
                "rows": len(s1),
                "sources": [relative(KPFM_H5), relative(PLATE_MAP_CSV)],
                "machine_readable_output": relative(TABLE_S1_CSV),
                "scope": (
                    "Paired acquisition order, plate/well identity, nominal composition, HDF5 acquisition IDs, "
                    "pixel dimensions, and manuscript-stated ternary scan area. Exact timestamps and external "
                    "calibration were not available and were not inferred."
                ),
            },
            "tables": {
                "S1": {
                    "rows": len(s1),
                    "sources": [relative(KPFM_H5), relative(PLATE_MAP_CSV)],
                    "machine_readable_output": relative(TABLE_S1_CSV),
                },
                "S2": {"rows": len(s2), "source": relative(SUMMARY_CSV)},
                "S3": {"rows": len(s3), "source": relative(RELATIONSHIPS_CSV)},
                "S4": {"rows": len(s4), "source": relative(ASSOCIATIONS_CSV)},
            },
            "source_sha256": {
                relative(path): sha256(path)
                for path in sources
            },
            "output": output.name,
            "output_sha256": sha256(temporary),
            "table_s1_csv_sha256": sha256(temporary_s1_csv),
            "checks": {
                "retained_relationships": 15,
                "unique_relationship_ids": 15,
                "photokpfm_acquisition_pairs": 19,
                "photokpfm_unique_wells": 13,
                "scan_pixels": "128 × 128",
                "ternary_scan_area": "5 × 5 µm²",
                "association_tests": 36,
                "bh_significant_tests": 0,
                "causal_claims_supported": 0,
            },
        }
        temporary_provenance = Path(directory) / provenance.name
        temporary_provenance.write_text(
            json.dumps(output_record, indent=2, ensure_ascii=False) + "\n",
            encoding="utf-8",
        )
        os.replace(temporary, output)
        os.replace(temporary_s1_csv, TABLE_S1_CSV)
        os.replace(temporary_provenance, provenance)
    return output, provenance, TABLE_S1_CSV


def main() -> None:
    args = parse_args()
    output, provenance, table_s1_csv = write_outputs(args.output.resolve(), args.overwrite)
    print(f"Generated: {output}")
    print(f"Table S1 CSV: {table_s1_csv}")
    print(f"Provenance: {provenance}")
    print("Status: complete; Tables S1-S4 were generated from release-local sources.")


if __name__ == "__main__":
    main()
